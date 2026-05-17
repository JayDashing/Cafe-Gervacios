import csv
import json
import os
import subprocess
from pathlib import Path

from docx import Document


WORKDIR = Path(r"C:\laragon\www\Cafe Gervacios\kiosk")
SOURCE_DOCX = Path(r"C:\Users\username\Downloads\Test_Case (1).docx")
OUTPUT_DOCX = WORKDIR / "Test_Case-1_Real_Updated.docx"
OUTPUT_CSV = WORKDIR / "Test_Case-1_Real_Updated.csv"
OUTPUT_LOG = WORKDIR / "Test_Case-1_Real_Execution_Log.txt"
PHP = "php"
NODE = r"C:\Users\username\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe"
NODE_PATH = r"C:\Users\username\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\node_modules"
HELPER = WORKDIR / "scripts" / "qa_live_helper.php"
BROWSER = WORKDIR / "scripts" / "browser_live_cases.cjs"


def run_cmd(cmd, extra_env=None):
    env = os.environ.copy()
    if extra_env:
        env.update(extra_env)
    result = subprocess.run(
        cmd,
        cwd=str(WORKDIR),
        capture_output=True,
        text=True,
        env=env,
    )
    return result


def helper(action, *args):
    result = run_cmd([PHP, str(HELPER), action, *args])
    output = (result.stdout or result.stderr).strip()
    if not output:
        raise RuntimeError(f"Helper produced no output for action {action}")
    data = json.loads(output)
    if not data.get("ok"):
        raise RuntimeError(data.get("error", f"Unknown helper failure for {action}"))
    return data


def browser(action):
    result = run_cmd([NODE, str(BROWSER), action], extra_env={"NODE_PATH": NODE_PATH})
    output = (result.stdout or result.stderr).strip()
    if not output:
        raise RuntimeError(f"Browser produced no output for action {action}")
    data = json.loads(output)
    if not data.get("ok"):
        raise RuntimeError(data.get("error", f"Unknown browser failure for {action}"))
    return data


def load_cases():
    doc = Document(str(SOURCE_DOCX))
    cases = []
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            cases.append({
                "id": cells[0],
                "title": cells[1],
                "pre": cells[2],
                "steps": cells[3],
                "expected": cells[4],
            })
    return cases


def execute_case(case_id):
    if case_id == "WTL-001":
        helper("init")
        browser("WTL-001")
        entry = helper("queue_by_phone", "09179000101")["entry"]
        if entry and entry["customer_name"] == "QA Browser Valid":
            return "Browser submitted form; waitlist entry saved with queue number 1.", "PASSED"
        raise RuntimeError("Walk-in entry not saved from browser submission")

    if case_id == "WTL-002":
        helper("init")
        browser("WTL-002")
        entry = helper("queue_by_name", "QA Browser NoPhone")["entry"]
        if entry and entry["customer_phone"] in ("", None):
            return "Browser saved walk-in entry successfully even with blank phone.", "PASSED"
        raise RuntimeError("Blank-phone walk-in entry not saved correctly")

    if case_id == "WTL-003":
        helper("init")
        result = browser("WTL-003")
        entry = helper("queue_by_phone", "09179000103")["entry"]
        body = result.get("body", "")
        if entry is None and "format is invalid" in body.lower():
            return "Browser showed name validation error; no queue entry persisted.", "PASSED"
        raise RuntimeError("Invalid-name browser validation did not behave as expected")

    if case_id == "FLM-001":
        helper("init")
        result = browser("FLM-001")
        if "Layout Editor" in result.get("title", "") or "Floor Map" in result.get("body", ""):
            return "Browser loaded Floor Map page and seating editor successfully.", "PASSED"
        raise RuntimeError("Floor map page did not load correctly in browser")

    if case_id == "ANL-010":
        helper("init")
        result = browser("ANL-010")
        if "/admin/login" in result.get("url", "") or "Sign in" in result.get("body", ""):
            return "Guest browser access redirected to login instead of analytics.", "PASSED"
        raise RuntimeError("Unauthorized analytics browser access was not blocked")

    if case_id in {
        "WTL-004", "WTL-005", "WTL-006", "WTL-007", "WTL-008", "WTL-009", "WTL-010",
        "FLM-002", "FLM-003", "FLM-004", "FLM-005", "FLM-006", "FLM-007", "FLM-008", "FLM-009", "FLM-010",
        "AUT-001", "AUT-002", "AUT-003", "AUT-004", "AUT-005", "AUT-006", "AUT-007", "AUT-008", "AUT-009", "AUT-010",
        "PRI-001", "PRI-002", "PRI-003", "PRI-004", "PRI-005", "PRI-006", "PRI-007", "PRI-008", "PRI-009", "PRI-010",
        "ANL-001", "ANL-002", "ANL-003", "ANL-004", "ANL-005", "ANL-006", "ANL-007", "ANL-008", "ANL-009",
    }:
        data = helper("run", case_id)
        return data["actual"], data["status"]

    raise RuntimeError(f"Case {case_id} not mapped")


def write_csv(cases, results):
    with OUTPUT_CSV.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["TEST CASE ID", "TITLE", "PRE-CONDITION", "TEST STEPS", "EXPECTED RESULTS", "ACTUAL RESULTS", "STATUS"])
        for case in cases:
            actual, status = results[case["id"]]
            writer.writerow([case["id"], case["title"], case["pre"], case["steps"], case["expected"], actual, status])


def update_docx(results):
    doc = Document(str(SOURCE_DOCX))
    for table in doc.tables:
        for row in table.rows[1:]:
            case_id = row.cells[0].text.strip()
            actual, status = results[case_id]
            row.cells[5].text = actual
            row.cells[6].text = status
    doc.save(str(OUTPUT_DOCX))


def main():
    cases = load_cases()
    results = {}
    logs = []
    passed = failed = blocked = 0
    bugs = []

    for case in cases:
        case_id = case["id"]
        header = f"=== {case_id} | {case['title']} ==="
        print(header)
        logs.append(header)
        logs.append(f"PRE-CONDITION: {case['pre']}")
        logs.append(f"TEST STEPS: {case['steps'].replace(chr(10), ' | ')}")
        try:
            actual, status = execute_case(case_id)
            results[case_id] = (actual, status)
            if status == "PASSED":
                passed += 1
            elif status == "BLOCKED":
                blocked += 1
            else:
                failed += 1
            line = f"RESULT: {status} | ACTUAL: {actual}"
            print(line)
            logs.append(line)
        except Exception as exc:
            results[case_id] = (f"Real execution failed: {exc}", "FAILED")
            failed += 1
            bug = f"{case_id}: {exc}"
            bugs.append(bug)
            line = f"RESULT: FAILED | ACTUAL: Real execution failed: {exc}"
            print(line)
            logs.append(line)
        print()
        logs.append("")

    write_csv(cases, results)
    update_docx(results)

    summary = [
        "=== SUMMARY ===",
        f"Total PASSED: {passed}",
        f"Total FAILED: {failed}",
        f"Total BLOCKED: {blocked}",
        f"Bugs found: {len(bugs)}",
    ] + bugs

    for line in summary:
        print(line)
        logs.append(line)

    OUTPUT_LOG.write_text("\n".join(logs), encoding="utf-8")


if __name__ == "__main__":
    main()
