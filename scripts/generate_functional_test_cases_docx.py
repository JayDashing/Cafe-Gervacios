import csv

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt


TEST_CASES = [
    # Subsystem 1 - Waitlist
    {
        "id": "WTL-001",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Create waitlist entry with valid walk-in details",
        "pre": "A logged-in admin or staff user is on the waitlist screen.",
        "steps": [
            "Enter a valid customer name.",
            "Enter a valid phone number.",
            "Enter a valid party size.",
            "Select a valid priority type.",
            "Submit the walk-in registration form.",
        ],
        "expected": "The system persists the new waitlist entry, assigns a queue number, calculates estimated wait, and shows the new record in the queue.",
    },
    {
        "id": "WTL-002",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Create waitlist entry when optional phone is blank",
        "pre": "A logged-in admin or staff user is on the waitlist screen.",
        "steps": [
            "Enter a valid customer name.",
            "Leave the phone number field blank.",
            "Enter a valid party size.",
            "Select a valid priority type.",
            "Submit the walk-in registration form.",
        ],
        "expected": "The system persists the new waitlist entry without a phone number and keeps the remaining guest details intact.",
    },
    {
        "id": "WTL-003",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Read waitlist entry using name or phone search",
        "pre": "At least one waitlist entry exists with a known customer name and phone number.",
        "steps": [
            "Search using the guest name.",
            "Confirm the matching queue record is returned.",
            "Search using the guest phone number.",
            "Confirm the same queue record is returned.",
        ],
        "expected": "The system returns the correct persisted waitlist record with the expected guest name, phone, party size, priority, and queue status.",
    },
    {
        "id": "WTL-004",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Update waitlist guest details before seating",
        "pre": "A waiting waitlist entry exists and the user has permission to edit queue records.",
        "steps": [
            "Open the existing waitlist entry for editing.",
            "Change the guest name, party size, and priority type.",
            "Save the changes.",
            "Reload or search for the same entry.",
        ],
        "expected": "The system updates the persisted waitlist entry and the refreshed queue shows the edited guest name, party size, priority, and recalculated queue order when applicable.",
    },
    {
        "id": "WTL-005",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Update waiting guest to notified when table-ready SMS is sent",
        "pre": "A waiting guest exists on the waitlist and the user has permission to update the entry.",
        "steps": [
            "Open the waitlist panel.",
            "Trigger the manual SMS action for the waiting guest.",
        ],
        "expected": "The guest status changes to notified, a hold expiration is set, a hold confirmation code is generated, and the table-ready SMS job is dispatched.",
    },
    {
        "id": "WTL-006",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Update notified guest to seated with correct hold code",
        "pre": "A notified guest exists with a valid hold confirmation code and a compatible table is selected.",
        "steps": [
            "Enter the correct hold confirmation code for the notified guest.",
            "Trigger the confirm-and-seat action using a compatible table.",
        ],
        "expected": "The system seats the guest successfully, updates the queue entry status, and marks the selected table as occupied.",
    },
    {
        "id": "WTL-007",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Reject waitlist update when hold confirmation code is incorrect",
        "pre": "A notified guest exists with a stored hold confirmation code and a compatible table is selected.",
        "steps": [
            "Enter an incorrect hold confirmation code.",
            "Trigger the confirm-and-seat action.",
        ],
        "expected": "The system does not seat the guest and shows an incorrect confirmation code error.",
    },
    {
        "id": "WTL-008",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Delete waitlist entry by cancelling held guest",
        "pre": "A waitlist entry exists with a reserved table hold and the user has permission to delete the entry.",
        "steps": [
            "Trigger the cancel action for the selected waitlist entry.",
        ],
        "expected": "The system marks the waitlist entry as cancelled and releases the reserved table back to available status.",
    },
    {
        "id": "WTL-009",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Update notified hold expiration by five minutes",
        "pre": "A waitlist entry exists in notified status with a current hold expiration.",
        "steps": [
            "Trigger the extend hold action for the notified entry.",
        ],
        "expected": "The system extends the hold expiration time by five minutes and shows a success message.",
    },
    {
        "id": "WTL-010",
        "subsystem": "Subsystem 1 - Waitlist",
        "title": "Read queue order after create, update, and delete actions",
        "pre": "Multiple waitlist entries exist with different priority types and at least one entry has been cancelled.",
        "steps": [
            "Create a new regular waitlist entry.",
            "Update another entry to a higher priority type.",
            "Cancel one existing entry.",
            "Refresh the queue list.",
            "Compare the visible queue order against active, non-cancelled records sorted by priority and arrival time.",
        ],
        "expected": "The queue list excludes deleted/cancelled entries, includes newly created records, applies priority updates, and displays records in the correct persisted order.",
    },

    # Subsystem 2 - Floor Map
    {
        "id": "FLM-001",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Read saved floor map table and seat records",
        "pre": "A table exists with multiple mapped seats, capacity, status, and floor coordinates.",
        "steps": [
            "Log in as an authorized admin user.",
            "Open the seating layout page.",
            "Read the loaded layout payload or rendered editor data for the seeded table.",
            "Compare the loaded label, capacity, status, seat count, and coordinates against the saved database records.",
        ],
        "expected": "The editor hydrates the exact saved table and seat data, preserving coordinates and allowing the user to target the existing table for later edit actions.",
    },
    {
        "id": "FLM-002",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Create new table on floor map with valid position",
        "pre": "An authenticated admin user is on the seating layout page.",
        "steps": [
            "Submit a create-seat placement request with valid position coordinates and table details.",
        ],
        "expected": "The system creates a new table and seat marker and returns the new layout information.",
    },
    {
        "id": "FLM-003",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Reject create table request with invalid coordinates",
        "pre": "An authenticated admin user is on the seating layout page.",
        "steps": [
            "Submit a placement request using invalid position coordinates outside the accepted range.",
        ],
        "expected": "The system rejects the placement request and shows coordinate validation errors.",
    },
    {
        "id": "FLM-004",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Update existing table label and capacity",
        "pre": "An authenticated admin user is editing an existing floor map table.",
        "steps": [
            "Submit an update request with a new valid label and capacity for an existing mapped table.",
        ],
        "expected": "The system updates the table metadata and returns the refreshed table and seat details.",
    },
    {
        "id": "FLM-005",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Reject update when capacity is below mapped seat count",
        "pre": "An authenticated admin user is editing a table that has multiple mapped seats.",
        "steps": [
            "Submit a table capacity update using a value lower than the current number of mapped seats.",
        ],
        "expected": "The system rejects the update and shows that capacity must be at least equal to the number of seats on the map.",
    },
    {
        "id": "FLM-006",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Create grouped table from multiple seat markers",
        "pre": "An authenticated admin user is on the seating layout page and multiple valid seat markers exist.",
        "steps": [
            "Select multiple existing seat markers.",
            "Submit the group action with a valid table label or default label generation.",
        ],
        "expected": "The system creates a grouped table, reassigns the selected seat markers to it, and preserves their floor positions.",
    },
    {
        "id": "FLM-007",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Delete single seat from floor map table",
        "pre": "An authenticated admin user is editing a table that has more than one mapped seat and no bookings.",
        "steps": [
            "Trigger the delete action for one seat using seat scope.",
        ],
        "expected": "The system removes the selected seat, resequences remaining seat indexes, and updates table capacity accordingly.",
    },
    {
        "id": "FLM-008",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Delete entire floor map table with no bookings",
        "pre": "An authenticated admin user is editing a table that has no bookings.",
        "steps": [
            "Trigger the delete action using table scope for the selected table.",
        ],
        "expected": "The system removes the entire table from the floor map and returns the removed table identifier.",
    },
    {
        "id": "FLM-009",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Reject delete when floor map table has bookings",
        "pre": "An authenticated admin user is editing a table that has one or more bookings.",
        "steps": [
            "Attempt to delete the table or one of its seats from the floor map.",
        ],
        "expected": "The system blocks the delete action and reports that tables with bookings cannot be removed.",
    },
    {
        "id": "FLM-010",
        "subsystem": "Subsystem 2 - Floor Map",
        "title": "Update dashboard seat action mode and clear stale selection",
        "pre": "An authenticated staff or admin user is on the dashboard seat map and a table is currently selected.",
        "steps": [
            "Select a table on the dashboard seat map.",
            "Change the seat click mode to waitlist.",
            "Attempt the next click using the waitlist mode.",
            "Change the seat click mode to table.",
            "Verify that the previous selected table is cleared and the next click follows the new mode.",
        ],
        "expected": "The system stores the new click mode, clears stale table selection, and routes the next seat-map click through the active mode instead of the old selection.",
    },

    # Subsystem 3 - Automation
    {
        "id": "AUT-001",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Update expired queue hold and create automation log",
        "pre": "A queue entry exists in notified status with a hold expiration time in the past and queue hold automation is enabled.",
        "steps": [
            "Run the queue hold automation task.",
        ],
        "expected": "The system cancels the notified hold, releases the reserved table, dispatches the queue-skipped SMS job, and records the automation log.",
    },
    {
        "id": "AUT-002",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Update wait estimate and create extended-wait alert",
        "pre": "Wait estimate automation is enabled and a waiting queue entry exists whose recalculated wait exceeds the alert threshold increase.",
        "steps": [
            "Run the wait estimates automation task.",
        ],
        "expected": "The system recalculates estimated wait, updates the queue entry, dispatches a wait-extended SMS job when the threshold is met, and records the automation log.",
    },
    {
        "id": "AUT-003",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Update overdue booking to no-show through automation",
        "pre": "No-show automation is enabled and an eligible booking exists past the configured no-show threshold.",
        "steps": [
            "Run the no-show automation task.",
        ],
        "expected": "The system marks the booking as cancelled with a no-show timestamp, releases any reserved table when applicable, dispatches the no-show SMS job, and triggers next queue notification.",
    },
    {
        "id": "AUT-004",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Update booking with late check-in SMS timestamp",
        "pre": "Late check-in automation is enabled and an eligible booking exists beyond the configured late check-in threshold without check-in.",
        "steps": [
            "Run the late check-in automation task.",
        ],
        "expected": "The system dispatches the late check-in SMS job, stores the late-check-in sent timestamp, and records the automation log.",
    },
    {
        "id": "AUT-005",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Update booking with 24-hour reminder timestamp",
        "pre": "Reminder automation is enabled and a paid booking exists within the 24-hour reminder window.",
        "steps": [
            "Run the reminders automation task.",
        ],
        "expected": "The system dispatches the 24-hour reminder SMS job, stores the 24-hour reminder timestamp, and records the automation log.",
    },
    {
        "id": "AUT-006",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Update booking with 2-hour reminder timestamp",
        "pre": "Reminder automation is enabled and a paid booking exists within the 2-hour reminder window.",
        "steps": [
            "Run the reminders automation task.",
        ],
        "expected": "The system dispatches the 2-hour reminder SMS job, stores the 2-hour reminder timestamp, and records the automation log.",
    },
    {
        "id": "AUT-007",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Delete booking-table assignment for cancelled or failed booking",
        "pre": "A booking exists with a linked reserved table and the booking has status cancelled or payment status failed.",
        "steps": [
            "Run the reservation table release automation task.",
        ],
        "expected": "The system releases the reserved table, deletes or clears the booking table assignment, and records the automation log.",
    },
    {
        "id": "AUT-008",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Read dedicated queue-hold setting before updating expired hold",
        "pre": "Master automation is disabled and a notified queue hold has already expired.",
        "steps": [
            "Run the queue hold automation task.",
        ],
        "expected": "The queue hold expiry still executes according to its dedicated gating rule even while master automation is disabled.",
    },
    {
        "id": "AUT-009",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Read disabled master setting and skip general update",
        "pre": "Master automation is disabled and a booking or queue state exists that would otherwise qualify for a general automation task.",
        "steps": [
            "Run a general automation task such as reminders or no-shows.",
        ],
        "expected": "The system skips the task and does not perform the general automation action while master automation is disabled.",
    },
    {
        "id": "AUT-010",
        "subsystem": "Subsystem 3 - Automation",
        "title": "Create failure log and admin alert when automation fails",
        "pre": "Automation admin alerts are enabled, an admin alert phone is configured, and an automation task encounters an internal failure.",
        "steps": [
            "Trigger an automation task failure condition.",
        ],
        "expected": "The system records the failure, logs the automation error, and dispatches an automation error SMS job to the configured admin alert phone.",
    },

    # Subsystem 4 - Priority
    {
        "id": "PRI-001",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Create PWD queue entry with correct priority score",
        "pre": "The priority service is available.",
        "steps": [
            "Create or evaluate a queue entry with priority type set to PWD.",
        ],
        "expected": "The system saves or evaluates the queue entry with priority score 100.",
    },
    {
        "id": "PRI-002",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Create pregnant queue entry with correct priority score",
        "pre": "The priority service is available.",
        "steps": [
            "Create or evaluate a queue entry with priority type set to pregnant.",
        ],
        "expected": "The system saves or evaluates the queue entry with priority score 100.",
    },
    {
        "id": "PRI-003",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Create senior queue entry with correct priority score",
        "pre": "The priority service is available.",
        "steps": [
            "Create or evaluate a queue entry with priority type set to senior.",
        ],
        "expected": "The system saves or evaluates the queue entry with priority score 100.",
    },
    {
        "id": "PRI-004",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Create regular queue entry with zero priority score",
        "pre": "The priority service is available.",
        "steps": [
            "Create or evaluate a queue entry with priority type set to none.",
        ],
        "expected": "The system saves or evaluates the queue entry with priority score 0.",
    },
    {
        "id": "PRI-005",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Read priority queue order after creating mixed entries",
        "pre": "The waitlist can accept one regular queue entry and one priority queue entry.",
        "steps": [
            "Create a regular queue entry with priority type none.",
            "Create a second queue entry with priority type senior, PWD, or pregnant.",
            "Request the ordered waitlist collection used by the panel.",
            "Compare the position and priority score of both entries.",
        ],
        "expected": "The priority entry receives a higher score and appears ahead of the regular entry in the ordered waitlist collection.",
    },
    {
        "id": "PRI-006",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Read PWD accessible-table rule before table assignment",
        "pre": "The setting queue_pwd_requires_accessible_table is enabled and a PWD queue entry exists.",
        "steps": [
            "Attempt to match or notify the PWD queue entry using a non-accessible table.",
        ],
        "expected": "The system does not treat the non-accessible table as a valid match for the PWD guest.",
    },
    {
        "id": "PRI-007",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Update PWD accessible-table setting and allow standard table",
        "pre": "The setting queue_pwd_requires_accessible_table is disabled and a PWD queue entry exists.",
        "steps": [
            "Attempt to match or seat the PWD queue entry using a compatible standard table.",
        ],
        "expected": "The system allows the standard table to be considered a valid fit for the PWD guest.",
    },
    {
        "id": "PRI-008",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Update senior queue entry to seated at compatible table",
        "pre": "A senior-priority queue entry exists and a compatible non-accessible table is available.",
        "steps": [
            "Attempt to match or seat the senior guest using the non-accessible compatible table.",
        ],
        "expected": "The system allows the senior guest to use the compatible table because accessible-only restriction applies only to PWD when configured.",
    },
    {
        "id": "PRI-009",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Update pregnant queue entry to seated at compatible table",
        "pre": "A pregnant-priority queue entry exists and a compatible non-accessible table is available.",
        "steps": [
            "Attempt to match or seat the pregnant guest using the non-accessible compatible table.",
        ],
        "expected": "The system allows the pregnant guest to use the compatible table because accessible-only restriction does not apply.",
    },
    {
        "id": "PRI-010",
        "subsystem": "Subsystem 4 - Priority",
        "title": "Create priority audit log when priority guest is seated",
        "pre": "A priority queue entry exists and is seated successfully at a valid table.",
        "steps": [
            "Seat the priority guest at the selected table.",
        ],
        "expected": "The system records a priority seating audit log entry containing the queue entry, table, accessibility, and wait-time details.",
    },

    # Subsystem 5 - Analytics
    {
        "id": "ANL-001",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read analytics count after creating current-day booking",
        "pre": "No QA bookings exist for the current day and the analytics component can query booking records.",
        "steps": [
            "Create one paid active booking with booked_at set to the current day.",
            "Open or refresh the seating analytics data source.",
            "Read the Bookings today metric.",
            "Compare the metric against the current-day booking query count.",
        ],
        "expected": "The Bookings today metric increases by one and matches the database count for bookings whose booked_at date is today.",
    },
    {
        "id": "ANL-002",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read analytics count after creating future booking",
        "pre": "The system has one booking scheduled today and one booking scheduled tomorrow.",
        "steps": [
            "Create a paid active booking with booked_at set to tomorrow.",
            "Refresh the seating analytics data source.",
            "Read the Bookings today metric.",
            "Compare the metric with a database query filtered to today's date only.",
        ],
        "expected": "The future booking is not counted in Bookings today; only records whose booked_at date is today are included.",
    },
    {
        "id": "ANL-003",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read checked-in metric after updating booking check-in",
        "pre": "A paid active booking exists for today with checked_in_at empty.",
        "steps": [
            "Set checked_in_at on the booking to the current timestamp.",
            "Refresh the seating analytics data source.",
            "Read the Checked in metric.",
            "Compare the metric against the count of bookings with checked_in_at today.",
        ],
        "expected": "The Checked in metric increases and matches the database count of bookings checked in today.",
    },
    {
        "id": "ANL-004",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read queue-seated metric after updating waitlist status",
        "pre": "A waiting queue entry and an available compatible table exist.",
        "steps": [
            "Seat the queue entry at the compatible table.",
            "Refresh the seating analytics data source.",
            "Read the Seated from queue metric.",
            "Compare the metric against queue entries with seated_at today.",
        ],
        "expected": "The Seated from queue metric increases and matches the number of queue entries seated today.",
    },
    {
        "id": "ANL-005",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read table counts after updating table status",
        "pre": "One QA table is available and one QA table is occupied.",
        "steps": [
            "Read the initial Free / occupied table metric.",
            "Change the available table to occupied.",
            "Refresh the seating analytics data source.",
            "Compare free and occupied counts against the current table status records.",
        ],
        "expected": "The free count decreases and the occupied count increases according to the current table status values in the database.",
    },
    {
        "id": "ANL-006",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read peak-hour bucket after creating booking",
        "pre": "No QA booking exists for the selected test hour within the last 7 days.",
        "steps": [
            "Create an active paid booking with booked_at set to a known hour today.",
            "Refresh the seating analytics chart payload.",
            "Read the peak-hour data array for that hour.",
            "Compare the bucket value against bookings for that hour in the last 7 days.",
        ],
        "expected": "Only the matching hour bucket increases, and the peak-hour dataset still contains exactly 24 hourly buckets.",
    },
    {
        "id": "ANL-007",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read table usage ranking after creating booking history",
        "pre": "Two or more tables exist and booking history can be created for the last 30 days.",
        "steps": [
            "Create two paid active bookings assigned to Table A.",
            "Create one paid active booking assigned to Table B.",
            "Refresh the top-table analytics payload.",
            "Read the ordered top-table list.",
        ],
        "expected": "Table A appears above Table B with the correct usage count, and the list remains capped to the top five tables.",
    },
    {
        "id": "ANL-008",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read zeroed analytics after deleting booking and queue data",
        "pre": "QA booking, queue, and table records can be safely removed from the test database.",
        "steps": [
            "Delete all QA bookings and queue entries.",
            "Refresh the analytics data source.",
            "Read booking, checked-in, seated, peak-hour, and top-table metrics.",
        ],
        "expected": "Count metrics return zero, peak-hour buckets return zero values, and top-table usage returns an empty list without runtime errors.",
    },
    {
        "id": "ANL-009",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Reject unauthorized read of analytics data",
        "pre": "A guest user is not authenticated and analytics data exists in the database.",
        "steps": [
            "Navigate directly to the seating analytics URL as a guest.",
            "Attempt to retrieve the analytics page or component data.",
            "Observe the response URL and content.",
        ],
        "expected": "The system redirects the guest to login or returns an authorization failure and does not expose analytics counts or chart payloads.",
    },
    {
        "id": "ANL-010",
        "subsystem": "Subsystem 5 - Analytics",
        "title": "Read peak-hour trend after deleting cancelled bookings from scope",
        "pre": "An active booking and a cancelled booking exist in the same hour within the last 7 days.",
        "steps": [
            "Refresh the seating analytics chart payload.",
            "Read the peak-hour bucket for the shared booking hour.",
            "Compare the bucket value against active and completed bookings only.",
        ],
        "expected": "The cancelled booking is excluded from the peak-hour trend while the active booking remains counted.",
    },
]


def format_steps(steps):
    return "\n".join(f"{index}. {step}" for index, step in enumerate(steps, start=1))


def format_cell(cell, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.bold = bold


def add_headers(table) -> None:
    headers = [
        "TEST CASE ID",
        "TITLE",
        "PRE-CONDITION",
        "TEST STEPS",
        "EXPECTED RESULTS",
        "ACTUAL RESULTS",
        "STATUS",
    ]
    widths = [0.85, 1.45, 1.6, 2.05, 2.0, 1.35, 0.9]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.text = header
        cell.width = Inches(width)
        format_cell(cell, bold=True)


def add_subsystem_row(table, subsystem: str) -> None:
    cells = table.add_row().cells
    cells[0].text = subsystem
    for index in range(1, 7):
        cells[0].merge(cells[index])
    format_cell(cells[0], bold=True)


def add_case_row(table, case) -> None:
    row = table.add_row().cells
    values = [
        case["id"],
        case["title"],
        case["pre"],
        format_steps(case["steps"]),
        case["expected"],
        "",
        "Not Executed",
    ]
    for cell, value in zip(row, values):
        cell.text = value
        format_cell(cell)


def add_case_table(document, cases) -> None:
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.autofit = False
    add_headers(table)

    current_subsystem = None
    for case in cases:
        if case["subsystem"] != current_subsystem:
            current_subsystem = case["subsystem"]
            add_subsystem_row(table, current_subsystem)
        add_case_row(table, case)


def chunk_cases(cases, chunk_size: int = 3):
    for index in range(0, len(cases), chunk_size):
        yield cases[index:index + chunk_size]


def build_docx(output_path: str) -> None:
    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(8)

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    document.add_heading("CRUD-Focused Functional Test Cases by Subsystem", level=0)
    document.add_paragraph(
        "These cases prioritize Create, Read, Update, and Delete behavior. Each case must prove persisted data, state change, validation, permissions, or report output after a real CRUD action."
    )

    for page_index, cases in enumerate(chunk_cases(TEST_CASES)):
        if page_index:
            document.add_page_break()
        add_case_table(document, cases)

    document.save(output_path)


def build_csv(output_path: str) -> None:
    headers = [
        "TEST CASE ID",
        "TITLE",
        "PRE-CONDITION",
        "TEST STEPS",
        "EXPECTED RESULTS",
        "ACTUAL RESULTS",
        "STATUS",
    ]
    with open(output_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(headers)
        for case in TEST_CASES:
            writer.writerow([
                case["id"],
                case["title"],
                case["pre"],
                format_steps(case["steps"]),
                case["expected"],
                "",
                "Not Executed",
            ])


if __name__ == "__main__":
    build_docx("C:\\laragon\\www\\Cafe Gervacios\\kiosk\\Functional_Test_Cases_By_Subsystem.docx")
    build_csv("C:\\laragon\\www\\Cafe Gervacios\\kiosk\\Functional_Test_Cases_By_Subsystem.csv")
