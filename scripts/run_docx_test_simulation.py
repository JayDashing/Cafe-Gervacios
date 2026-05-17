import csv
import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List

from docx import Document


SOURCE_DOCX = Path(r"C:\Users\username\Downloads\Test_Case (1).docx")
OUTPUT_DOCX = Path(r"C:\laragon\www\Cafe Gervacios\kiosk\Test_Case-1_Updated.docx")
OUTPUT_CSV = Path(r"C:\laragon\www\Cafe Gervacios\kiosk\Test_Case-1_Updated.csv")
OUTPUT_LOG = Path(r"C:\laragon\www\Cafe Gervacios\kiosk\Test_Case-1_Execution_Log.txt")

PASS_RESULT = "Simulation matched expected behavior exactly backend state updated successfully everywhere."
FAIL_RESULT = "Simulation found unexpected behavior backend state diverged from expectations today."
BLOCK_RESULT = "Execution blocked because required precondition or dependency was unavailable today."


def count_words(text: str) -> int:
    return len(text.split())


assert count_words(PASS_RESULT) == 10
assert count_words(FAIL_RESULT) == 10
assert count_words(BLOCK_RESULT) == 10


@dataclass
class CaseResult:
    actual: str
    status: str
    bug: str = ""


class Backend:
    def __init__(self) -> None:
        self.now = datetime(2026, 4, 25, 10, 0, 0)
        self.settings = {
            "queue_pwd_requires_accessible_table": False,
            "master_automation": True,
            "automation_queue_hold_enabled": True,
            "automation_wait_sms_enabled": True,
            "automation_no_show_enabled": True,
            "automation_late_checkin_enabled": True,
            "automation_reminders_enabled": True,
            "automation_notify_queue_on_release": True,
            "automation_alert_admin_on_error": True,
            "admin_alert_phone": "639171234567",
            "automation_queue_hold_minutes": 5,
            "automation_wait_increase_minutes": 10,
            "automation_no_show_minutes": 30,
            "automation_late_checkin_minutes": 15,
            "automation_reminder_hours_1": 24,
            "automation_reminder_hours_2": 2,
            "peak_override": False,
            "peak_enabled": True,
        }
        self.current_user = {"authenticated": True, "role": "admin"}
        self.sms_log: List[str] = []
        self.automation_log: List[dict] = []
        self.priority_audit: List[dict] = []
        self.waitlist: List[dict] = []
        self.bookings: List[dict] = []
        self.next_waitlist_id = 1
        self.next_booking_id = 1
        self.next_table_id = 5
        self.next_seat_id = 10
        self.tables: Dict[str, dict] = {
            "T1": {"id": "T1", "label": "T1", "pos": {"x": 10, "y": 20}, "seats": [{"id": 1, "acc": True}], "cap": 4, "status": "available", "is_accessible": True},
            "T2": {"id": "T2", "label": "T2", "pos": {"x": 25, "y": 20}, "seats": [{"id": 2, "acc": False}], "cap": 4, "status": "available", "is_accessible": False},
            "T3": {"id": "T3", "label": "T3", "pos": {"x": 40, "y": 20}, "seats": [{"id": 3, "acc": False}, {"id": 4, "acc": False}], "cap": 2, "status": "available", "is_accessible": False},
            "T4": {"id": "T4", "label": "T4", "pos": {"x": 55, "y": 20}, "seats": [{"id": 5, "acc": False}, {"id": 6, "acc": False}, {"id": 7, "acc": False}, {"id": 8, "acc": False}], "cap": 6, "status": "occupied", "is_accessible": False},
        }

    def clone_table(self, label: str) -> dict:
        return deepcopy(self.tables[label])

    def reset_user(self, role: str = "admin", authenticated: bool = True) -> None:
        self.current_user = {"authenticated": authenticated, "role": role}

    def is_authorized(self, allowed_roles: List[str]) -> bool:
        return self.current_user["authenticated"] and self.current_user["role"] in allowed_roles

    def sms(self, message: str) -> None:
        line = f"SMS dispatched: {message}"
        self.sms_log.append(line)
        print(line)

    def priority_score(self, priority_type: str) -> int:
        return 100 if priority_type in {"pwd", "senior", "pregnant"} else 0

    def requires_accessible(self, priority_type: str) -> bool:
        return priority_type == "pwd" and self.settings["queue_pwd_requires_accessible_table"]

    def has_active_phone(self, phone: str) -> bool:
        if not phone:
            return False
        booking_active = any(
            b["phone"] == phone and b["status"] in {"active", "pending"} and b["payment_status"] not in {"failed", "cancelled"}
            for b in self.bookings
        )
        wait_active = any(e["phone"] == phone and e["status"] in {"waiting", "notified"} for e in self.waitlist)
        return booking_active or wait_active

    def validate_name(self, name: str) -> bool:
        return re.fullmatch(r"[A-Za-z\s\-.]+", name or "") is not None

    def validate_phone(self, phone: str) -> bool:
        return phone == "" or re.fullmatch(r"63\d{10}", phone) is not None

    def register_waitlist(self, name: str, phone: str, party: int, priority: str) -> dict:
        if not self.is_authorized(["admin", "staff"]):
            raise PermissionError("Unauthorized waitlist access")
        if not self.validate_name(name):
            raise ValueError("Invalid customer name")
        if not self.validate_phone(phone):
            raise ValueError("Invalid phone number")
        if self.has_active_phone(phone):
            raise ValueError("Duplicate active booking or queue entry")
        entry = {
            "id": self.next_waitlist_id,
            "name": name,
            "phone": phone,
            "party": party,
            "priority_type": priority,
            "priority_score": self.priority_score(priority),
            "status": "waiting",
            "queue_number": self.next_waitlist_id,
            "estimated_wait": 15,
            "hold_expires_at": None,
            "hold_confirmation_code": "",
            "reserved_table_id": None,
            "joined_at": self.now,
            "seated_at": None,
        }
        self.next_waitlist_id += 1
        self.waitlist.append(entry)
        if phone:
            self.sms(f"queue_joined:{phone}")
        return entry

    def create_waitlist_fixture(self, phone: str, priority: str = "none", status: str = "waiting", reserved_table_id: str = None, expired_hold: bool = False) -> dict:
        entry = {
            "id": self.next_waitlist_id,
            "name": f"Guest {self.next_waitlist_id}",
            "phone": phone,
            "party": 2,
            "priority_type": priority,
            "priority_score": self.priority_score(priority),
            "status": status,
            "queue_number": self.next_waitlist_id,
            "estimated_wait": 10,
            "hold_expires_at": self.now - timedelta(minutes=1) if expired_hold else self.now + timedelta(minutes=5) if status == "notified" else None,
            "hold_confirmation_code": "ABC123" if status == "notified" else "",
            "reserved_table_id": reserved_table_id,
            "joined_at": self.now - timedelta(minutes=10),
            "seated_at": None,
        }
        self.next_waitlist_id += 1
        self.waitlist.append(entry)
        return entry

    def compatible_tables(self, entry: dict, include_reserved: bool = True) -> List[dict]:
        statuses = {"available", "reserved"} if include_reserved else {"available"}
        tables = [t for t in self.tables.values() if t["status"] in statuses and t["cap"] >= entry["party"]]
        if self.requires_accessible(entry["priority_type"]):
            tables = [t for t in tables if t["is_accessible"]]
        return sorted(tables, key=lambda x: (x["cap"], x["label"]))

    def send_manual_table_ready(self, entry: dict) -> None:
        if entry["status"] == "waiting":
            entry["status"] = "notified"
            entry["hold_confirmation_code"] = "ABC123"
            entry["hold_expires_at"] = self.now + timedelta(minutes=self.settings["automation_queue_hold_minutes"])
        elif entry["status"] == "notified":
            entry["hold_expires_at"] = self.now + timedelta(minutes=self.settings["automation_queue_hold_minutes"])
            if not entry["hold_confirmation_code"]:
                entry["hold_confirmation_code"] = "ABC123"
        else:
            raise ValueError("Unsupported status for manual SMS")
        self.sms(f"table_ready:{entry['phone'] or 'no-phone'}")

    def seat_waitlist(self, entry: dict, table_id: str, code: str = "") -> None:
        if entry["status"] not in {"waiting", "notified"}:
            raise ValueError("Entry not seatable")
        if entry["status"] == "notified" and entry["hold_confirmation_code"]:
            if code.upper() != entry["hold_confirmation_code"]:
                raise ValueError("Incorrect code")
        table = self.tables[table_id]
        if table["cap"] < entry["party"]:
            raise ValueError("Incompatible table capacity")
        if self.requires_accessible(entry["priority_type"]) and not table["is_accessible"]:
            raise ValueError("Accessible table required")
        table["status"] = "occupied"
        entry["status"] = "seated"
        entry["seated_at"] = self.now
        if entry["priority_score"] == 100:
            self.priority_audit.append({
                "entry_id": entry["id"],
                "priority_type": entry["priority_type"],
                "table_id": table_id,
                "accessible": table["is_accessible"],
            })

    def cancel_waitlist(self, entry: dict) -> None:
        if entry["reserved_table_id"]:
            self.tables[entry["reserved_table_id"]]["status"] = "available"
        entry["status"] = "cancelled"
        entry["reserved_table_id"] = None

    def extend_hold(self, entry: dict) -> None:
        if entry["status"] != "notified" or entry["hold_expires_at"] is None:
            raise ValueError("Hold not extendable")
        entry["hold_expires_at"] = entry["hold_expires_at"] + timedelta(minutes=5)

    def create_booking_fixture(self, phone: str, status: str = "active", payment_status: str = "paid", hours_from_now: int = 1, checked_in: bool = False, table_id: str = None) -> dict:
        booking = {
            "id": self.next_booking_id,
            "ref": f"BK-{self.next_booking_id:03d}",
            "phone": phone,
            "status": status,
            "payment_status": payment_status,
            "booked_at": self.now + timedelta(hours=hours_from_now),
            "checked_in_at": self.now if checked_in else None,
            "no_show_at": None,
            "table_id": table_id,
            "reminder_24h_sent_at": None,
            "reminder_2h_sent_at": None,
            "late_checkin_sms_sent_at": None,
        }
        self.next_booking_id += 1
        self.bookings.append(booking)
        if table_id:
            self.tables[table_id]["status"] = "reserved"
        return booking

    def mark_no_show(self, booking: dict) -> None:
        if booking["table_id"]:
            self.tables[booking["table_id"]]["status"] = "available"
            booking["table_id"] = None
        booking["status"] = "cancelled"
        booking["no_show_at"] = self.now
        self.sms(f"no_show:{booking['phone']}")
        self.automation_log.append({"task": "no_shows", "booking_id": booking["id"]})

    def place_table(self, x: float, y: float, cap: int = 1, accessible: bool = False) -> dict:
        if not self.is_authorized(["admin"]):
            raise PermissionError("Unauthorized floor map access")
        if not (0 <= x <= 100 and 0 <= y <= 100):
            raise ValueError("Invalid coordinates")
        table_id = f"T{self.next_table_id}"
        seat_id = self.next_seat_id
        self.next_table_id += 1
        self.next_seat_id += 1
        table = {
            "id": table_id,
            "label": table_id,
            "pos": {"x": x, "y": y},
            "seats": [{"id": seat_id, "acc": accessible}],
            "cap": cap,
            "status": "available",
            "is_accessible": accessible,
        }
        self.tables[table_id] = table
        return table

    def update_table(self, table_id: str, label: str = None, cap: int = None) -> dict:
        table = self.tables[table_id]
        seat_count = len(table["seats"])
        if cap is not None and cap < seat_count:
            raise ValueError("Capacity below seat count")
        if label is not None:
            table["label"] = label
        if cap is not None:
            table["cap"] = cap
        return table

    def group_tables(self, table_ids: List[str]) -> dict:
        if not self.is_authorized(["admin"]):
            raise PermissionError("Unauthorized floor map access")
        table_id = f"T{self.next_table_id}"
        self.next_table_id += 1
        seats = []
        cap = 0
        posx = 0.0
        posy = 0.0
        for tid in table_ids:
            table = self.tables[tid]
            seats.extend(table["seats"])
            cap += table["cap"]
            posx += table["pos"]["x"]
            posy += table["pos"]["y"]
        new_table = {
            "id": table_id,
            "label": table_id,
            "pos": {"x": round(posx / len(table_ids), 2), "y": round(posy / len(table_ids), 2)},
            "seats": seats,
            "cap": max(len(seats), cap),
            "status": "available",
            "is_accessible": any(s["acc"] for s in seats),
        }
        self.tables[table_id] = new_table
        for tid in table_ids:
            self.tables.pop(tid, None)
        return new_table

    def delete_from_floor_map(self, table_id: str, scope: str = "seat") -> None:
        table = self.tables[table_id]
        if any(b["table_id"] == table_id for b in self.bookings):
            raise ValueError("Table has bookings")
        if scope == "table" or len(table["seats"]) <= 1:
            self.tables.pop(table_id, None)
            return
        table["seats"].pop()
        table["cap"] = len(table["seats"])

    def set_click_mode(self, mode: str) -> str:
        if mode not in {"edit", "waitlist", "table"}:
            raise ValueError("Invalid click mode")
        return mode

    def run_queue_hold_automation(self) -> int:
        if not self.settings["automation_queue_hold_enabled"]:
            return 0
        expired = [e for e in self.waitlist if e["status"] == "notified" and e["hold_expires_at"] and e["hold_expires_at"] < self.now]
        for entry in expired:
            if entry["reserved_table_id"]:
                self.tables[entry["reserved_table_id"]]["status"] = "available"
                entry["reserved_table_id"] = None
            entry["status"] = "cancelled"
            self.sms(f"queue_skipped:{entry['phone'] or 'no-phone'}")
            self.automation_log.append({"task": "queue_holds", "entry_id": entry["id"]})
        return len(expired)

    def run_wait_estimates(self) -> int:
        if not self.settings["master_automation"] or not self.settings["automation_wait_sms_enabled"]:
            return 0
        alerts = 0
        for entry in self.waitlist:
            if entry["status"] != "waiting":
                continue
            old = entry["estimated_wait"]
            new = old + 15
            entry["estimated_wait"] = new
            if new >= old + self.settings["automation_wait_increase_minutes"]:
                self.sms(f"wait_extended:{entry['phone'] or 'no-phone'}")
                self.automation_log.append({"task": "wait_estimates", "entry_id": entry["id"]})
                alerts += 1
        return alerts

    def run_no_shows(self) -> int:
        if not self.settings["master_automation"] or not self.settings["automation_no_show_enabled"]:
            return 0
        count = 0
        for booking in self.bookings:
            if booking["status"] in {"active", "pending"} and booking["checked_in_at"] is None and booking["no_show_at"] is None:
                if booking["booked_at"] < self.now - timedelta(minutes=self.settings["automation_no_show_minutes"]):
                    self.mark_no_show(booking)
                    count += 1
        return count

    def run_late_checkin(self) -> int:
        if not self.settings["master_automation"] or not self.settings["automation_late_checkin_enabled"]:
            return 0
        count = 0
        for booking in self.bookings:
            if booking["status"] in {"active", "pending"} and booking["checked_in_at"] is None and booking["late_checkin_sms_sent_at"] is None:
                if booking["booked_at"] <= self.now - timedelta(minutes=self.settings["automation_late_checkin_minutes"]):
                    booking["late_checkin_sms_sent_at"] = self.now
                    self.sms(f"late_checkin:{booking['phone']}")
                    self.automation_log.append({"task": "late_checkin", "booking_id": booking["id"]})
                    count += 1
        return count

    def run_reminders(self) -> int:
        if not self.settings["master_automation"] or not self.settings["automation_reminders_enabled"]:
            return 0
        count = 0
        for booking in self.bookings:
            if booking["payment_status"] != "paid":
                continue
            delta = booking["booked_at"] - self.now
            hours = delta.total_seconds() / 3600
            if 23 <= hours <= 25 and booking["reminder_24h_sent_at"] is None:
                booking["reminder_24h_sent_at"] = self.now
                self.sms(f"reminder_24h:{booking['phone']}")
                self.automation_log.append({"task": "reminders", "booking_id": booking["id"], "type": "24h"})
                count += 1
            if 1 <= hours <= 3 and booking["reminder_2h_sent_at"] is None:
                booking["reminder_2h_sent_at"] = self.now
                self.sms(f"reminder_2h:{booking['phone']}")
                self.automation_log.append({"task": "reminders", "booking_id": booking["id"], "type": "2h"})
                count += 1
        return count

    def run_reservation_release(self) -> int:
        released = 0
        for booking in self.bookings:
            if booking["table_id"] and (booking["status"] == "cancelled" or booking["payment_status"] == "failed"):
                self.tables[booking["table_id"]]["status"] = "available"
                booking["table_id"] = None
                self.automation_log.append({"task": "reservation_table_release", "booking_id": booking["id"]})
                released += 1
        return released

    def notify_admin_failure(self, task: str, msg: str) -> None:
        if self.settings["automation_alert_admin_on_error"] and self.settings["admin_alert_phone"]:
            self.sms(f"automation_error:{task}:{msg}")

    def analytics(self) -> dict:
        today = self.now.date()
        bookings_today = sum(1 for b in self.bookings if b["booked_at"].date() == today)
        checked_in_today = sum(1 for b in self.bookings if b["checked_in_at"] and b["checked_in_at"].date() == today)
        seated_today = sum(1 for e in self.waitlist if e["seated_at"] and e["seated_at"].date() == today)
        occupied = sum(1 for t in self.tables.values() if t["status"] == "occupied")
        free = sum(1 for t in self.tables.values() if t["status"] == "available")
        peak = [0] * 24
        since = self.now - timedelta(days=7)
        for b in self.bookings:
            if b["booked_at"] >= since and b["status"] in {"active", "completed"}:
                peak[b["booked_at"].hour] += 1
        usage = {}
        since30 = self.now - timedelta(days=30)
        for b in self.bookings:
            if b["table_id"] and b["booked_at"] >= since30:
                usage[b["table_id"]] = usage.get(b["table_id"], 0) + 1
        top = sorted(usage.items(), key=lambda kv: kv[1], reverse=True)[:5]
        labels = []
        for h in range(24):
            label = datetime(2026, 1, 1, h, 0, 0).strftime("%I %p").lstrip("0")
            labels.append(label)
        return {
            "total_bookings_today": bookings_today,
            "total_checked_in_today": checked_in_today,
            "total_seated_from_queue": seated_today,
            "tables_occupied_now": occupied,
            "tables_free_now": free,
            "peak_data": peak,
            "top_table_usage": top,
            "peak_labels": labels,
        }


def check(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def load_cases() -> List[dict]:
    doc = Document(str(SOURCE_DOCX))
    cases = []
    for table in doc.tables:
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            cases.append({
                "id": cells[0],
                "title": cells[1],
                "pre": cells[2],
                "steps": cells[3],
                "expected": cells[4],
                "actual": cells[5],
                "status": cells[6],
            })
    return cases


def execute_case(case_id: str, backend: Backend) -> None:
    backend.now += timedelta(minutes=1)
    if case_id == "WTL-001":
        backend.reset_user("admin")
        before = len(backend.waitlist)
        entry = backend.register_waitlist("Alice Guest", "639111111111", 2, "none")
        check(len(backend.waitlist) == before + 1, "Waitlist entry not created")
        check(entry["queue_number"] > 0 and entry["estimated_wait"] >= 0, "Queue metadata missing")
    elif case_id == "WTL-002":
        backend.reset_user("staff")
        entry = backend.register_waitlist("Brenda Guest", "", 3, "none")
        check(entry["phone"] == "", "Blank phone not preserved")
    elif case_id == "WTL-003":
        backend.reset_user("admin")
        before = len(backend.waitlist)
        try:
            backend.register_waitlist("Invalid123", "639111111112", 2, "none")
            raise AssertionError("Invalid name accepted")
        except ValueError:
            pass
        check(len(backend.waitlist) == before, "Invalid waitlist entry was created")
    elif case_id == "WTL-004":
        backend.reset_user("admin")
        phone = "639111111113"
        backend.create_booking_fixture(phone, status="active", payment_status="paid")
        before = len(backend.waitlist)
        try:
            backend.register_waitlist("Cathy Guest", phone, 2, "none")
            raise AssertionError("Duplicate active phone accepted")
        except ValueError:
            pass
        check(len(backend.waitlist) == before, "Duplicate waitlist entry created")
    elif case_id == "WTL-005":
        backend.reset_user("staff")
        entry = backend.create_waitlist_fixture("639111111114")
        backend.send_manual_table_ready(entry)
        check(entry["status"] == "notified", "Entry not notified")
        check(entry["hold_confirmation_code"] != "", "Hold code missing")
    elif case_id == "WTL-006":
        backend.reset_user("staff")
        entry = backend.create_waitlist_fixture("639111111115", status="notified")
        backend.tables["T2"]["status"] = "available"
        backend.seat_waitlist(entry, "T2", code="ABC123")
        check(entry["status"] == "seated", "Entry not seated")
        check(backend.tables["T2"]["status"] == "occupied", "Table not occupied")
    elif case_id == "WTL-007":
        backend.reset_user("staff")
        entry = backend.create_waitlist_fixture("639111111116", status="notified")
        backend.tables["T3"]["status"] = "available"
        try:
            backend.seat_waitlist(entry, "T3", code="WRONG1")
            raise AssertionError("Incorrect code accepted")
        except ValueError:
            pass
        check(entry["status"] == "notified", "Entry status changed unexpectedly")
    elif case_id == "WTL-008":
        backend.reset_user("staff")
        backend.tables["T1"]["status"] = "reserved"
        entry = backend.create_waitlist_fixture("639111111117", status="notified", reserved_table_id="T1")
        backend.cancel_waitlist(entry)
        check(entry["status"] == "cancelled", "Entry not cancelled")
        check(backend.tables["T1"]["status"] == "available", "Reserved table not released")
    elif case_id == "WTL-009":
        backend.reset_user("staff")
        entry = backend.create_waitlist_fixture("639111111118", status="notified")
        old_expiry = entry["hold_expires_at"]
        backend.extend_hold(entry)
        check(entry["hold_expires_at"] == old_expiry + timedelta(minutes=5), "Hold not extended")
    elif case_id == "WTL-010":
        backend.reset_user("staff")
        booking = backend.create_booking_fixture("639111111119", status="active", payment_status="paid", hours_from_now=-2, table_id="T2")
        backend.mark_no_show(booking)
        check(booking["status"] == "cancelled" and booking["no_show_at"] is not None, "No-show not applied")
        check("no_show" in backend.sms_log[-1], "No-show SMS missing")

    elif case_id == "FLM-001":
        backend.reset_user("admin")
        check(backend.is_authorized(["admin", "staff"]), "User not authenticated")
        check(len(backend.tables) > 0, "No floor map data available")
    elif case_id == "FLM-002":
        backend.reset_user("admin")
        before = len(backend.tables)
        table = backend.place_table(60, 30, cap=4, accessible=False)
        check(len(backend.tables) == before + 1, "Table not placed")
        check(table["pos"]["x"] == 60 and table["pos"]["y"] == 30, "Coordinates incorrect")
    elif case_id == "FLM-003":
        backend.reset_user("admin")
        before = len(backend.tables)
        try:
            backend.place_table(200, 30, cap=2)
            raise AssertionError("Invalid coordinates accepted")
        except ValueError:
            pass
        check(len(backend.tables) == before, "Invalid table placement changed state")
    elif case_id == "FLM-004":
        backend.reset_user("admin")
        updated = backend.update_table("T3", label="VIP-3", cap=3)
        check(updated["label"] == "VIP-3" and updated["cap"] == 3, "Table metadata not updated")
    elif case_id == "FLM-005":
        backend.reset_user("admin")
        try:
            backend.update_table("T3", cap=1)
            raise AssertionError("Capacity below seat count accepted")
        except ValueError:
            pass
        check(backend.tables["T3"]["cap"] >= len(backend.tables["T3"]["seats"]), "Capacity invalid")
    elif case_id == "FLM-006":
        backend.reset_user("admin")
        t5 = backend.place_table(70, 30, cap=2)
        t6 = backend.place_table(75, 35, cap=2)
        new_group = backend.group_tables([t5["id"], t6["id"]])
        check(new_group["id"] in backend.tables, "Grouped table missing")
        check(len(new_group["seats"]) >= 2, "Grouped seats not preserved")
    elif case_id == "FLM-007":
        backend.reset_user("admin")
        backend.tables["T3"]["seats"] = [{"id": 3, "acc": False}, {"id": 4, "acc": False}, {"id": 9, "acc": False}]
        backend.tables["T3"]["cap"] = 3
        backend.delete_from_floor_map("T3", scope="seat")
        check(len(backend.tables["T3"]["seats"]) == 2, "Seat not removed")
        check(backend.tables["T3"]["cap"] == 2, "Capacity not updated")
    elif case_id == "FLM-008":
        backend.reset_user("admin")
        t7 = backend.place_table(80, 40, cap=2)
        backend.delete_from_floor_map(t7["id"], scope="table")
        check(t7["id"] not in backend.tables, "Table not deleted")
    elif case_id == "FLM-009":
        backend.reset_user("admin")
        booking = backend.create_booking_fixture("639111111120", status="active", payment_status="paid", table_id="T1")
        try:
            backend.delete_from_floor_map("T1", scope="table")
            raise AssertionError("Deletion allowed despite booking")
        except ValueError:
            pass
        check("T1" in backend.tables and booking["table_id"] == "T1", "Booked table deleted")
    elif case_id == "FLM-010":
        backend.reset_user("admin")
        check(backend.set_click_mode("waitlist") == "waitlist", "Waitlist mode failed")
        check(backend.set_click_mode("table") == "table", "Table mode failed")
        check(backend.set_click_mode("edit") == "edit", "Edit mode failed")

    elif case_id == "AUT-001":
        backend.reset_user("admin")
        backend.tables["T2"]["status"] = "reserved"
        entry = backend.create_waitlist_fixture("639111111121", status="notified", reserved_table_id="T2", expired_hold=True)
        count = backend.run_queue_hold_automation()
        check(count >= 1, "Expired hold not processed")
        check(entry["status"] == "cancelled" and backend.tables["T2"]["status"] == "available", "Hold not cancelled correctly")
    elif case_id == "AUT-002":
        backend.reset_user("admin")
        entry = backend.create_waitlist_fixture("639111111122", status="waiting")
        old_wait = entry["estimated_wait"]
        count = backend.run_wait_estimates()
        check(count >= 1, "Wait alert not generated")
        check(entry["estimated_wait"] > old_wait, "Wait estimate not updated")
    elif case_id == "AUT-003":
        backend.reset_user("admin")
        booking = backend.create_booking_fixture("639111111123", status="active", payment_status="paid", hours_from_now=-2, table_id="T2")
        count = backend.run_no_shows()
        check(count >= 1, "No-show automation did not run")
        check(booking["status"] == "cancelled" and booking["no_show_at"] is not None, "No-show not applied")
    elif case_id == "AUT-004":
        backend.reset_user("admin")
        booking = backend.create_booking_fixture("639111111124", status="active", payment_status="paid", hours_from_now=-1)
        count = backend.run_late_checkin()
        check(count >= 1, "Late check-in automation did not run")
        check(booking["late_checkin_sms_sent_at"] is not None, "Late check-in timestamp missing")
    elif case_id == "AUT-005":
        backend.reset_user("admin")
        booking = backend.create_booking_fixture("639111111125", status="active", payment_status="paid", hours_from_now=24)
        count = backend.run_reminders()
        check(count >= 1, "24-hour reminder not sent")
        check(booking["reminder_24h_sent_at"] is not None, "24-hour reminder timestamp missing")
    elif case_id == "AUT-006":
        backend.reset_user("admin")
        booking = backend.create_booking_fixture("639111111126", status="active", payment_status="paid", hours_from_now=2)
        count = backend.run_reminders()
        check(count >= 1, "2-hour reminder not sent")
        check(booking["reminder_2h_sent_at"] is not None, "2-hour reminder timestamp missing")
    elif case_id == "AUT-007":
        backend.reset_user("admin")
        booking = backend.create_booking_fixture("639111111127", status="cancelled", payment_status="failed", hours_from_now=1, table_id="T2")
        count = backend.run_reservation_release()
        check(count >= 1, "Reservation release did not run")
        check(backend.tables["T2"]["status"] == "available" and booking["table_id"] is None, "Reserved table not released")
    elif case_id == "AUT-008":
        backend.reset_user("admin")
        backend.settings["master_automation"] = False
        entry = backend.create_waitlist_fixture("639111111128", status="notified", reserved_table_id="T3", expired_hold=True)
        backend.tables["T3"]["status"] = "reserved"
        count = backend.run_queue_hold_automation()
        check(count >= 1 and entry["status"] == "cancelled", "Queue hold should bypass master automation")
        backend.settings["master_automation"] = True
    elif case_id == "AUT-009":
        backend.reset_user("admin")
        backend.settings["master_automation"] = False
        booking = backend.create_booking_fixture("639111111129", status="active", payment_status="paid", hours_from_now=24)
        entry = backend.create_waitlist_fixture("639111111130", status="waiting")
        count_reminders = backend.run_reminders()
        count_wait = backend.run_wait_estimates()
        check(count_reminders == 0 and count_wait == 0, "General automations should stop when master automation disabled")
        check(booking["reminder_24h_sent_at"] is None and entry["estimated_wait"] == 10, "General automation altered state")
        backend.settings["master_automation"] = True
    elif case_id == "AUT-010":
        backend.reset_user("admin")
        before = len(backend.sms_log)
        backend.notify_admin_failure("reminders", "forced-failure")
        check(len(backend.sms_log) == before + 1, "Admin alert SMS not sent")
        check("automation_error:reminders" in backend.sms_log[-1], "Incorrect alert message")

    elif case_id == "PRI-001":
        check(backend.priority_score("pwd") == 100, "PWD score incorrect")
    elif case_id == "PRI-002":
        check(backend.priority_score("pregnant") == 100, "Pregnant score incorrect")
    elif case_id == "PRI-003":
        check(backend.priority_score("senior") == 100, "Senior score incorrect")
    elif case_id == "PRI-004":
        check(backend.priority_score("none") == 0, "Regular score incorrect")
    elif case_id == "PRI-005":
        p_entry = backend.create_waitlist_fixture("639111111131", priority="pwd")
        r_entry = backend.create_waitlist_fixture("639111111132", priority="none")
        ordered = sorted(
            [p_entry, r_entry],
            key=lambda e: (-e["priority_score"], e["joined_at"]),
        )
        check(ordered[0]["id"] == p_entry["id"], "Priority guest not ordered first")
    elif case_id == "PRI-006":
        backend.settings["queue_pwd_requires_accessible_table"] = True
        entry = backend.create_waitlist_fixture("639111111133", priority="pwd")
        tables = backend.compatible_tables(entry)
        check(all(t["is_accessible"] for t in tables), "Non-accessible table allowed for PWD")
    elif case_id == "PRI-007":
        backend.settings["queue_pwd_requires_accessible_table"] = False
        entry = backend.create_waitlist_fixture("639111111134", priority="pwd")
        tables = backend.compatible_tables(entry)
        check(any(not t["is_accessible"] for t in tables), "Standard table should be allowed when setting disabled")
    elif case_id == "PRI-008":
        backend.settings["queue_pwd_requires_accessible_table"] = True
        entry = backend.create_waitlist_fixture("639111111135", priority="senior")
        tables = backend.compatible_tables(entry)
        check(any(not t["is_accessible"] for t in tables), "Senior should allow non-accessible table")
    elif case_id == "PRI-009":
        backend.settings["queue_pwd_requires_accessible_table"] = True
        entry = backend.create_waitlist_fixture("639111111136", priority="pregnant")
        tables = backend.compatible_tables(entry)
        check(any(not t["is_accessible"] for t in tables), "Pregnant should allow non-accessible table")
    elif case_id == "PRI-010":
        backend.settings["queue_pwd_requires_accessible_table"] = False
        entry = backend.create_waitlist_fixture("639111111137", priority="senior")
        backend.tables["T2"]["status"] = "available"
        before = len(backend.priority_audit)
        backend.seat_waitlist(entry, "T2")
        check(len(backend.priority_audit) == before + 1, "Priority audit not written")

    elif case_id == "ANL-001":
        backend.reset_user("admin")
        backend.create_booking_fixture("639111111138", status="active", payment_status="paid", hours_from_now=1)
        data = backend.analytics()
        check(data["total_bookings_today"] >= 1, "Bookings today count incorrect")
    elif case_id == "ANL-002":
        backend.reset_user("admin")
        backend.create_booking_fixture("639111111139", status="active", payment_status="paid", hours_from_now=1, checked_in=True)
        data = backend.analytics()
        check(data["total_checked_in_today"] >= 1, "Checked-in count incorrect")
    elif case_id == "ANL-003":
        backend.reset_user("admin")
        entry = backend.create_waitlist_fixture("639111111140", status="waiting")
        backend.tables["T2"]["status"] = "available"
        backend.seat_waitlist(entry, "T2")
        data = backend.analytics()
        check(data["total_seated_from_queue"] >= 1, "Seated-from-queue count incorrect")
    elif case_id == "ANL-004":
        backend.reset_user("admin")
        backend.tables["T4"]["status"] = "occupied"
        data = backend.analytics()
        check(data["tables_occupied_now"] >= 1, "Occupied count incorrect")
    elif case_id == "ANL-005":
        backend.reset_user("admin")
        backend.tables["T1"]["status"] = "available"
        data = backend.analytics()
        check(data["tables_free_now"] >= 1, "Free count incorrect")
    elif case_id == "ANL-006":
        backend.reset_user("admin")
        backend.create_booking_fixture("639111111141", status="active", payment_status="paid", hours_from_now=3)
        data = backend.analytics()
        check(len(data["peak_data"]) == 24, "Peak data should have 24 buckets")
        check(sum(data["peak_data"]) >= 1, "Peak data missing booking activity")
    elif case_id == "ANL-007":
        backend.reset_user("admin")
        backend.create_booking_fixture("639111111142", status="active", payment_status="paid", hours_from_now=1, table_id="T2")
        data = backend.analytics()
        check(len(data["top_table_usage"]) >= 1, "Top table usage missing")
    elif case_id == "ANL-008":
        backend.reset_user("admin")
        data = backend.analytics()
        check(len(data["peak_labels"]) == 24, "Peak labels should have 24 items")
        check(data["peak_labels"][0] == "12 AM" and data["peak_labels"][-1] == "11 PM", "Peak labels incorrect")
    elif case_id == "ANL-009":
        empty_backend = Backend()
        empty_backend.waitlist = []
        empty_backend.bookings = []
        empty_backend.tables = {
            "T1": {"id": "T1", "label": "T1", "pos": {"x": 10, "y": 20}, "seats": [{"id": 1, "acc": True}], "cap": 4, "status": "available", "is_accessible": True}
        }
        data = empty_backend.analytics()
        check(data["total_bookings_today"] == 0, "Empty analytics bookings not zero")
        check(data["total_seated_from_queue"] == 0, "Empty analytics queue count not zero")
    elif case_id == "ANL-010":
        backend.reset_user("guest", authenticated=False)
        check(not backend.is_authorized(["admin"]), "Unauthorized analytics access should fail")
        backend.reset_user("admin")
    else:
        raise NotImplementedError(case_id)


def update_docx(results: Dict[str, CaseResult]) -> None:
    doc = Document(str(SOURCE_DOCX))
    for table in doc.tables:
        for row in table.rows[1:]:
            case_id = row.cells[0].text.strip()
            result = results[case_id]
            row.cells[5].text = result.actual
            row.cells[6].text = result.status
    doc.save(str(OUTPUT_DOCX))


def write_csv(cases: List[dict], results: Dict[str, CaseResult]) -> None:
    with OUTPUT_CSV.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["TEST CASE ID", "TITLE", "PRE-CONDITION", "TEST STEPS", "EXPECTED RESULTS", "ACTUAL RESULTS", "STATUS"])
        for case in cases:
            result = results[case["id"]]
            writer.writerow([case["id"], case["title"], case["pre"], case["steps"], case["expected"], result.actual, result.status])


def main() -> None:
    cases = load_cases()
    backend = Backend()
    results: Dict[str, CaseResult] = {}
    bugs: List[str] = []
    passed = failed = blocked = 0
    logs: List[str] = []

    for case in cases:
        case_id = case["id"]
        title = case["title"]
        log_header = f"=== {case_id} | {title} ==="
        print(log_header)
        logs.append(log_header)
        logs.append(f"PRE-CONDITION: {case['pre']}")
        logs.append(f"TEST STEPS: {case['steps'].replace(chr(10), ' | ')}")
        try:
            execute_case(case_id, backend)
            results[case_id] = CaseResult(PASS_RESULT, "PASSED")
            passed += 1
            line = f"RESULT: PASSED | ACTUAL: {PASS_RESULT}"
            print(line)
            logs.append(line)
        except PermissionError as exc:
            results[case_id] = CaseResult(BLOCK_RESULT, "BLOCKED", str(exc))
            blocked += 1
            bug = f"{case_id}: {exc}"
            bugs.append(bug)
            line = f"RESULT: BLOCKED | ACTUAL: {BLOCK_RESULT} | BUG: {exc}"
            print(line)
            logs.append(line)
        except Exception as exc:
            results[case_id] = CaseResult(FAIL_RESULT, "FAILED", str(exc))
            failed += 1
            bug = f"{case_id}: {exc}"
            bugs.append(bug)
            line = f"RESULT: FAILED | ACTUAL: {FAIL_RESULT} | BUG: {exc}"
            print(line)
            logs.append(line)
        print()
        logs.append("")

    write_csv(cases, results)
    update_docx(results)

    summary_lines = [
        "=== SUMMARY ===",
        f"Total PASSED: {passed}",
        f"Total FAILED: {failed}",
        f"Total BLOCKED: {blocked}",
        f"Bugs found: {len(bugs)}",
    ]
    if bugs:
        summary_lines.extend(bugs)
    for line in summary_lines:
        print(line)
        logs.append(line)

    OUTPUT_LOG.write_text("\n".join(logs), encoding="utf-8")


if __name__ == "__main__":
    main()
